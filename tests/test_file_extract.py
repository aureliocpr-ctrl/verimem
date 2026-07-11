"""File -> text extraction for the document RAG tier (roadmap #1: whole-file ingest).

Extract text from real file formats (txt/md, pdf, docx, html) so it can be
chunked (chunk_text) and embedded. The last test wires extract -> chunk to prove
the #1 pipeline end (provenance preserved).
"""
from __future__ import annotations

import pytest

from engram.file_extract import extract_text


def test_txt_and_md(tmp_path) -> None:
    f = tmp_path / "doc.txt"
    f.write_text("Hello world.\nSecond line.", encoding="utf-8")
    assert "Hello world" in extract_text(f)

    m = tmp_path / "doc.md"
    m.write_text("# Title\n\nBody text here.", encoding="utf-8")
    assert "Body text here" in extract_text(m)


def test_docx(tmp_path) -> None:
    import docx

    d = docx.Document()
    d.add_paragraph("Alpha paragraph content.")
    d.add_paragraph("Beta paragraph content.")
    p = tmp_path / "doc.docx"
    d.save(str(p))

    out = extract_text(p)
    assert "Alpha paragraph content" in out
    assert "Beta paragraph content" in out


def test_docx_zip_bomb_is_refused(tmp_path, monkeypatch) -> None:
    """Sicurezza (zip-bomb, audit E2 2026-07-11): _extract_docx delega a
    python-docx che decomprime word/document.xml senza limite — un DOCX da 74KB
    estraeva 40MB (ratio 529x, OOM a scala). Il guard sulla dimensione DICHIARATA
    nella central directory rifiuta il file PRIMA che la libreria lo decomprima."""
    import docx

    import engram.file_extract as fe
    monkeypatch.setattr(fe, "_MAX_MEMBER_BYTES", 1_000_000)
    monkeypatch.setattr(fe, "_MAX_TOTAL_BYTES", 4_000_000)
    d = docx.Document()
    d.add_paragraph("A" * 3_000_000)  # word/document.xml > 1MB cap
    p = tmp_path / "bomb.docx"
    d.save(str(p))
    with pytest.raises(ValueError, match="zip-bomb guard"):
        fe.extract_text(p)


def test_pdf(tmp_path) -> None:
    import fitz  # PyMuPDF

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Pdf sample text here.")
    p = tmp_path / "doc.pdf"
    doc.save(str(p))
    doc.close()

    assert "Pdf sample text" in extract_text(p)


def test_unsupported_extension_raises(tmp_path) -> None:
    f = tmp_path / "doc.xyz"
    f.write_text("data", encoding="utf-8")
    with pytest.raises(ValueError):
        extract_text(f)


def test_missing_file_raises(tmp_path) -> None:
    with pytest.raises(FileNotFoundError):
        extract_text(tmp_path / "nope.txt")


def test_extract_feeds_chunker_with_provenance(tmp_path) -> None:
    # The #1 pipeline end: a real file -> text -> provenance-anchored chunks.
    from engram.chunking import chunk_text

    f = tmp_path / "big.txt"
    f.write_text("word " * 500, encoding="utf-8")
    text = extract_text(f)
    chunks = chunk_text(text, chunk_size=300, overlap=50)
    assert len(chunks) > 1
    assert text[chunks[0].start:chunks[0].end] == chunks[0].text
